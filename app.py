from flask import Flask, render_template, request, jsonify, redirect, url_for, session , flash
from dotenv import load_dotenv
from groq import Groq
import os
from PyPDF2 import PdfReader
from docx import Document
import pytesseract
from PIL import Image
import sqlite3
from werkzeug.security import generate_password_hash, check_password_hash
from flask import send_file
from reportlab.platypus import Preformatted, SimpleDocTemplate, Paragraph, Image as RLImage, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document as DocxDocument
from io import BytesIO

load_dotenv()
client = Groq(api_key=os.getenv("GROQ_API_KEY"))

app = Flask(__name__)
app.secret_key = os.getenv("SECRET_KEY")

# DATABASE CONNECTION
def get_db():
    conn = sqlite3.connect("users.db")
    conn.row_factory = sqlite3.Row
    return conn

# CREATE TABLE
def create_table():
    conn = get_db()
    conn.execute("""
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            full_name TEXT,
            email TEXT,
            username TEXT UNIQUE,
            password TEXT
        )
    """)
    conn.commit()
    conn.close()

create_table()

# HOME
@app.route("/")
def home():
    return render_template("index.html", mcqs="")

# SIGNUP
@app.route("/signup", methods=["GET", "POST"])
def signup():
    if request.method == "POST":
        full_name = request.form["name"]
        email = request.form["email"]
        username = request.form["username"]
        password = request.form["password"]

        hashed_password = generate_password_hash(password)

        try:
            conn = get_db()
            conn.execute(
                "INSERT INTO users ( full_name, email, username, password) VALUES (?, ?, ?, ?)",
                (full_name, email, username, hashed_password)
            )
            conn.commit()
            conn.close()

            flash("Signup successful! Please login.")
            return redirect(url_for("login"))

        except sqlite3.IntegrityError:
            flash("Username already exists!")

    return render_template("signup.html")

# LOGIN
@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        username = request.form["username"]
        password = request.form["password"]

        conn = get_db()
        user = conn.execute(
            "SELECT * FROM users WHERE username=?",
            (username,)
        ).fetchone()
        conn.close()

        if user and check_password_hash(user["password"], password):
            session["user"] = user["username"]
            next_page = session.pop("after_login", None)

            if next_page == "quiz":
                return redirect(url_for("index", mode="quiz"))
                
            return redirect(url_for("index", mode="quiz"))
        else: 
            flash("Invalid username or password") 
            
    return render_template("login.html")

# DASHBOARD
@app.route("/dashboard")
def dashboard():
    if "user" in session:
        return f"Welcome {session['user']} 🎉"
    return redirect(url_for("login"))

# LOGOUT
@app.route("/logout")
def logout():
    session.pop("user", None)
    return redirect(url_for("login"))

# MAIN LOGIC
@app.route("/index", methods=["GET", "POST"])
def index():
    mcqs = ""
    mode = request.args.get("mode") or request.form.get("mode", "practice")
    if request.method == "POST":
        difficulty = request.form.get("difficulty")
        count = request.form.get("count")
        mode = request.form.get("mode", "practice")   

        if mode == "quiz" and "user" not in session:
            session["after_login"] = "quiz"
            return redirect(url_for("login"))

        prev_score = request.form.get("score") 
        count = int(count) if count else 5

        # Adaptive Difficulty
        if difficulty == "Adaptive" and prev_score:
            try:
                prev_score = int(prev_score)
                if prev_score > count * 0.7:
                    difficulty = "Hard"
                elif prev_score < count * 0.4:
                    difficulty = "Easy"
                else:
                    difficulty = "Medium"
            except:
                difficulty = "Medium"

        topic = request.form.get("topic", "").strip()

        pdf_file = request.files.get("pdf_file")
        txt_file = request.files.get("txt_file")
        docx_file = request.files.get("docx_file")
        image_file = request.files.get("image_file")

        extracted_text = ""
        source_used = "Topic"

        # TXT
        if txt_file and txt_file.filename != "":
            extracted_text = txt_file.read().decode("utf-8", errors="ignore")
            source_used = "TXT File"

        # PDF
        elif pdf_file and pdf_file.filename != "":
            reader = PdfReader(pdf_file)
            pdf_text = ""
            for page in reader.pages:
                pdf_text += (page.extract_text() or "") + "\n"
            extracted_text = pdf_text
            source_used = "PDF File"

        # DOCX
        elif docx_file and docx_file.filename != "":
            doc = Document(docx_file)
            extracted_text = "\n".join([p.text for p in doc.paragraphs])

        # IMAGE
        elif image_file and image_file.filename != "":
            img = Image.open(image_file).convert("RGB")
            extracted_text = pytesseract.image_to_string(img, lang="eng")
            source_used = "Image (OCR)"

        extracted_text = extracted_text.strip()[:6000]

        level_instruction = ""
        if difficulty:
            level_instruction = f"""
            The MCQs must strictly follow this exam level: {difficulty} 
            - If GATE: Focus on conceptual, numerical, and application-based questions. 
            - If NET: Include theoretical, conceptual, and research-oriented questions. 
            - If UPSC/CGPSC: Focus on factual + analytical + current-affairs style. 
            - If Bloom's Taxonomy: 
                * Remembering: Direct facts 
                * Understanding: Concept clarity 
                * Applying: Problem-solving 
                * Analyzing: Comparison & reasoning 
                * Evaluating: Judgement-based 
                * Creating: Scenario-based 
            - Maintain the exact tone and difficulty of the selected level. """

        # 🔥 PRACTICE MODE
        if mode == "practice":
            if extracted_text:
                prompt = f"""
                Generate exactly {count} MCQs from the given text. 
                
                FORMAT:
                Q1) Question
                A) ...
                B) ...
                C) ...
                D) ...
                Answer: <A/B/C/D>
                
                Q2) Question
                A) ...
                B) ...
                C) ...
                D) ...
                Answer: <A/B/C/D>
                
                RULES:
                1. Only MCQs
                2. DO NOT add explanations
                3. MUST include correct answer for each question
                4. Keep clean format
                
                TEXT:
                {extracted_text}
                """
            else:
                prompt = f"""
               Generate exactly {count} MCQs on topic: {topic}

                FORMAT:
                Q1) Question
                A) ...
                B) ...
                C) ...
                D) ...
                Answer: <A/B/C/D>
                
                RULES:
                1. Only MCQs
                2. MUST include correct answers
                3. No explanations
                """
        # 🎯 QUIZ MODE
        else:
            if extracted_text:
                prompt = f"""
                You are an expert exam question setter. 
                Task: 
                Generate exactly {count} HIGH-QUALITY questions from the given text. 
                IMPORTANT: 
                Generate a MIX of different question types, not only MCQs.
                Include:
                1. MCQs (Multiple Choice Questions) 
                2. Fill in the blanks 
                3. Short answer questions 
                4. One word answer questions 
                5. Case study based questions 
                6. True/False 
                7. Assertion-Reason (if applicable) 
                {level_instruction} 
                FORMAT: 
                Q1) (MCQ) 
                Question... 
                A) ... 
                B) ... 
                C) ... 
                D) ...
                Answer: <A/B/C/D>
                Explanation: ... 
                Q2) (Fill in the Blank) 
                Question with ______
                Answer: ... 
                Q3) (Short Answer) 
                Question... 
                Answer: ... 
                Q4) (One Word)
                Question... 
                Answer: ... 
                Q5) (Case Study) 
                <Small paragraph>
                Questions: 
                a) ... 
                b) ... 
                Answers:
                a) ... 
                b) ... 
                Q6) (True/False) 
                Statement... 
                Answer: True/False 
                Q7) (Assertion-Reason)
                Assertion: ... 
                Reason: ...
                Options: 
                A) Both true 
                B) Both false 
                C) Assertion true, Reason false 
                D) Assertion false, Reason true 
                Answer: ...
                RULES: 
                1. Questions must be from the text only 
                2. Do not repeat questions 
                3. Do not add extra commentary 
                4. Keep language simple and exam-oriented 
                5. Add 1-2 trusted reference links 
                STRICT FORMATTING RULES: 
                1. DO NOT use ** or any markdown symbols 
                2. DO NOT add extra blank lines between questions 
                3. DO NOT mention marks like (5 marks), (2 marks), etc. 
                4. Keep everything in plain text only 
                5. Each question must start exactly like: Q1) (Type)
                6. Question should be in the same line, no unnecessary spacing 
                7. Do NOT add notes or instructions like [Note: ...] 
                TEXT: 
                {extracted_text} """
            else:
                prompt = f"""
                You are an expert exam question setter.
                Generate exactly {count} HIGH-QUALITY questions on this topic: {topic} 
                IMPORTANT: 
                Generate a MIX of different question types. Include: 
                - MCQs 
                - Fill in the blanks 
                - Short answer 
                - One word 
                - Case study 
                - True/False 
                - Assertion-Reason 
                {level_instruction} 
                Follow this format: 
                Q1) (MCQ) ... 
                Q2) (Fill in the Blank) ... 
                Q3) ... 
                Q4) (One Word) ... 
                Q5) (Case Study) ... 
                Q6) (True/False) ... 
                Q7) (Assertion-Reason) ...
                RULES: 
                1. Do not repeat questions 
                2. Keep exam-level quality 
                3. Keep answers accurate 
                4. Add 1-2 trusted reference links """

        # ✅ COMMON API CALL (for both modes)
        response = client.chat.completions.create(
            model="llama-3.1-8b-instant",
            messages=[{"role": "user", "content": prompt}]
        )

        mcqs = response.choices[0].message.content
        mcqs = f"Source Used: {source_used} | Mode: {mode} | Level: {difficulty}\n\n" + mcqs

    if mcqs:
        if mode == "practice":
            return render_template("PracticeMode.html", mcqs=mcqs)
        else:
            return render_template("result.html", mcqs=mcqs, mode=mode)

    return render_template("index.html", mcqs=mcqs)


@app.route("/chat", methods=["POST"])
def chat():
    data = request.get_json()

    user_message = data.get("message")
    context = data.get("context")
    lang = data.get("lang", "en")

    if "translate" in user_message.lower():
        prompt = f"""
        You are a STRICT translation engine.

        Translate the text EXACTLY into {"pure Hindi (Devanagari script only, NOT Hinglish)" if lang=="hi" else "English"}.
        
        ABSOLUTE RULES (DO NOT BREAK):
        1. ONLY translate the given text.
        2. DO NOT explain anything.
        3. DO NOT add examples.
        4. DO NOT add words like "beta", "namaste", etc.
        5. DO NOT change structure or format.
        6. DO NOT add Answer or Explanation if not present.
        7. DO NOT modify numbering (Q1, A, B, etc).

        OUTPUT MUST LOOK EXACTLY SAME FORMAT, ONLY LANGUAGE CHANGED.

        TEXT:
        {context}
        """

    elif lang == "hi":
        prompt = f"""
        You are a helpful teacher who explains in simple Hindi.

        Context:
        {context}

        Student Question:
        {user_message}

        Explain in simple Hindi (Hinglish allowed).
        Make it easy for Indian students.
        """
    else:
        prompt = f"""
        You are a helpful teacher.

        Context:
        {context}

        Student Question:
        {user_message}

        Explain clearly in simple English.

        Also include:
        - Why correct answer is right
        - Why other options are wrong
        - Short concept summary
        """

    response = client.chat.completions.create(
        model="llama-3.1-8b-instant",
        temperature=0,   # 🔥 forces strict translation (no creativity)
        messages=[{"role": "user", "content": prompt}]
    )

    reply = response.choices[0].message.content

    sources = []

    if context:
        if "PDF File" in context:
            sources.append("Uploaded PDF File")
        elif "TXT File" in context:
            sources.append("Uploaded TXT File")
        elif "Image (OCR)" in context:
            sources.append("Extracted from Image (OCR)")
        else:
            sources.append("Generated from provided content")

    if not sources:
        sources = [
            "https://en.wikipedia.org/wiki/" + user_message.replace(" ", "_"),
            "https://www.geeksforgeeks.org/" + user_message.replace(" ", "-").lower()
        ]

    return jsonify({
        "reply": reply,
        "sources": sources
    })

@app.route("/export/<format>", methods=["POST"])
def export_file(format):
    content = request.form.get("mcqs", "")

    logo_path = "static/logo.png"   # 👉 keep your logo here

    # ===== PDF EXPORT =====
    if format == "pdf":
        buffer = BytesIO()
    
        doc = SimpleDocTemplate(buffer, pagesize=None)
        styles = getSampleStyleSheet()
    
        elements = []
    
        logo_path = "static/logo.png"
    
        # SAFE LOGO HANDLING
        if os.path.exists(logo_path):
            try:
                elements.append(RLImage(logo_path, width=120, height=60))
                elements.append(Spacer(1, 10))
            except:
                pass
    
        # CLEAN TEXT SPLITTING (VERY IMPORTANT FIX)
        cleaned_lines = content.split("\n")
    
        for line in cleaned_lines:
            line = line.strip()
    
            # Skip empty lines
            if not line:
                continue
    
            # Escape problematic characters (ReportLab issue fix)
            line = line.replace("&", "and")
    
            elements.append(Preformatted(line, styles["Normal"]))
            elements.append(Spacer(1, 5))
    
        doc.build(elements)
    
        buffer.seek(0)
    
        return send_file(
            buffer,
            as_attachment=True,
            download_name="MCQs.pdf",
            mimetype="application/pdf"
        )
    # ===== DOCX EXPORT =====
    elif format == "docx":
        doc = DocxDocument()

        # Add logo
        if os.path.exists(logo_path):
            doc.add_picture(logo_path, width=None)

        doc.add_paragraph(content)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return send_file(
            buffer,
            as_attachment=True,
            download_name="MCQs.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    # ===== TEXT EXPORT =====
    elif format == "txt":
        buffer = BytesIO()

        # Add logo name as header
        text_content = "=== QUIZ BUILDER ===\n\n" + content

        buffer.write(text_content.encode("utf-8"))
        buffer.seek(0)

        return send_file(
            buffer,
            as_attachment=True,
            download_name="MCQs.txt",
            mimetype="text/plain"
        )

    return "Invalid format"

if __name__ == "__main__":
    app.run(debug=True)
